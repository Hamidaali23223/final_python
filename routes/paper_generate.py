from fastapi import APIRouter, FastAPI
from fastapi import Form
from pydantic import BaseModel
from typing import List, Dict
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from openai import OpenAI
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, json, hashlib, re
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
import requests
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
import chromadb
from config import env


router = APIRouter()

ENDPOINT=env.ENDPOINT
TOKEN=env.TOKEN
MODEL_NAME=env.MODEL_NAME

client = OpenAI(
    base_url=ENDPOINT,
    api_key=TOKEN,
)
class PaperRequest(BaseModel):
    university: str
    subject: str
    department: str
    semester: str
    term: str
    dflevelMCQ: str
    dflevelShort: str
    dflevelLong: str
    total_marks: int
    duration_minutes: int
    questions_count: Dict[str, int]

def generate_exam_paper(syllabus_text, count_dict, data, topic_query):

    question_types = ["MCQ", "Short", "Long"]
    paper = {}

    print("syllabustext", syllabus_text)
    print("count_dict", count_dict)
    print("data", data)
    print(data.dflevelMCQ)
    print(data.dflevelShort)
    print(data.dflevelLong)

    for qtype in question_types:

        if qtype in count_dict and count_dict[qtype] > 0:
            if qtype == "MCQ":
                prompt = (
        f"You are an AI exam paper generator. Generate exactly {count_dict[qtype]} "
        f"Topic: {topic_query}, Make sure you focus on relevant sources from {syllabus_text}"
        f"{data.dflevelMCQ.lower()} level MCQ questions based on this syllabus: {syllabus_text}\n\n"
        f"Format each MCQ exactly like this:\n"
        f"1. [Question text]\nA) [Option A]\nB) [Option B]\nC) [Option C]\nD) [Option D]\n\n"
        f"Rules:\n"
        f"- Options must be brief (1-2 words).\n"
        f"- Questions should match {data.dflevelMCQ.lower()} difficulty.\n"
        f"- Do NOT include answers or extra text.\n"
        f"- Only generate {count_dict[qtype]} questions — no more, no less."
    )
            elif qtype == "Short":
                prompt = (
        f"You are an AI exam paper generator. Generate exactly {count_dict[qtype]} "
        f"Topic: {topic_query}, Make sure you focus on relevant sources from {syllabus_text}"
        f"{data.dflevelShort.lower()} level short-answer questions from the syllabus: {syllabus_text}\n\n"
        f"Each question should be formatted like this:\n1. [Short question]\n\n"
        f"Rules:\n"
        f"- Question should be 1 lines.\n"
        f"- Reflect {data.dflevelShort.lower()} difficulty.\n"
        f"- No answers or extra info.\n"
        f"- Only generate {count_dict[qtype]} questions."
    )
            else:  # Long questions
                prompt = (
        f"You are an AI exam paper generator. Generate exactly {count_dict[qtype]} "
        f"Topic: {topic_query}, Make sure you focus on relevant sources from {syllabus_text}"
        f"{data.dflevelLong.lower()} level long-answer questions based on the syllabus: {syllabus_text}\n\n"
        f"Format like this:\n1. [Detailed question]\n\n"
        f"Rules:\n"
        f"- Each question must require a detailed explanation.\n"
        f"- Match {data.dflevelLong.lower()} difficulty.\n"
        f"- Keep question text concise (1-2 lines).\n"
        f"- No answers or extra commentary.\n"
        f"- Generate exactly {count_dict[qtype]} questions only."
    )

            # === GPT Call ===
            response = client.chat.completions.create(
                temperature=1.0,
                top_p=1.0,
                max_tokens=4000,
                model=MODEL_NAME,
                messages=[
                    {"role": "system", "content": "You are an exam paper generator AI. You follow instructions strictly."},
                    {"role": "user", "content": prompt}
                ]
            )

            content = response.choices[0].message.content.strip()
            print("Generated response:", response)

            # === Extract Questions ===
            if qtype == "MCQ":
                mcq_list = []
                pattern = r'(\d+)\.\s+(.*?)(?:\n|\r\n)A\)\s+(.*?)(?:\n|\r\n)B\)\s+(.*?)(?:\n|\r\n)C\)\s+(.*?)(?:\n|\r\n)D\)\s+(.*?)(?:\n\n|\r\n\r\n|$)'
                matches = re.findall(pattern, content, re.DOTALL)
                for match in matches:
                    mcq_list.append({
                        "question": match[1].strip(),
                        "options": [match[2].strip(), match[3].strip(), match[4].strip(), match[5].strip()]
                    })
                if not mcq_list:
                    mcq_list = [q.strip() for q in content.split('\n') if q.strip()]
                paper[qtype] = mcq_list[:count_dict[qtype]]
            else:
                matches = re.findall(r'(\d+)\.\s+(.*?)(?:\n\n|\r\n\r\n|$)', content, re.DOTALL)
                questions = [match[1].strip() for match in matches] if matches else [q.strip() for q in content.split('\n') if q.strip()]
                paper[qtype] = questions[:count_dict[qtype]]

    return paper

def format_mcq_options(doc, options):
    # Create a paragraph for all options, starting on a new line
    print("doc", doc)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0)
    
    # Set up tab stops for aligned options
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.75))  # Position for option B
    tab_stops.add_tab_stop(Inches(3.5))   # Position for option C
    tab_stops.add_tab_stop(Inches(5.25))  # Position for option D
    
    # Add options with proper tabs to ensure alignment
    option_letters = ['A)', 'B)', 'C)', 'D)']
    
    # Add first option without tab
    if options and len(options) > 0:
        run = p.add_run(f"{option_letters[0]} {options[0]}")
        
    # Add remaining options with tabs
    for i in range(1, min(4, len(options))):
        run = p.add_run(f"\t{option_letters[i]} {options[i]}")
    print("p",p)
    return p

def format_paper_pu(questions, data):
    doc = DocxDocument()

    # Set up document formatting
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Add header information
    title = doc.add_heading({data.university}, 0)
    title.alignment = 1  # Center alignment
    
    info = doc.add_paragraph()
    info.alignment = 1  # Center alignment
    depart = info.add_run(f"Department of {data.department}")
    depart.bold = True
    depart.font.size = Pt(16)

    subject = doc.add_paragraph()
    subject.alignment = 1  # Center alignment  
    subject.add_run(f"Semester: {data.semester} {data.term} Term 2025").bold = True          
    
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(3.5), WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7), WD_ALIGN_PARAGRAPH.RIGHT)

    p.add_run(f"Subject: {data.subject}\t")
    p.add_run(f"Marks: {data.total_marks}\t")
    p.add_run(f"Time: {data.duration_minutes} hrs").bold = True
    
    doc.add_paragraph("") 
    p = doc.add_paragraph()
    
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7), WD_ALIGN_PARAGRAPH.RIGHT)

    # Add left-aligned "Name" field
    left_run = p.add_run("Name: ____________________\t")
    left_run.bold = True

    # Add right-aligned "Roll Number" field 
    right_run = p.add_run("Roll Number: ________________")
    right_run.bold = True

    # Handle MCQs
    if "MCQ" in questions and questions["MCQ"]:
        doc.add_heading("Q.1 - Multiple Choice Questions", level=1)
        
        for i, q in enumerate(questions["MCQ"], 1):
            if isinstance(q, dict) and "question" in q and "options" in q:
                # Question on its own line
                p_question = doc.add_paragraph(f"{i}. {q['question']}")
                p_question.paragraph_format.space_after = Pt(0)  # Reduce space
                
                # Options on the next line, aligned
                format_mcq_options(doc, q["options"])
            else:
                # Handle raw string format
                q_text = q if isinstance(q, str) else str(q)
                if q_text.startswith('-'):
                    q_text = q_text[1:].strip()
                doc.add_paragraph(f"{i}. {q_text}")
            
            # Add small space after each complete MCQ
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
    # Handle Short Questions
    if "Short" in questions and questions["Short"]:
        doc.add_heading("Q.2 - Short Questions", level=1)
        for i, q in enumerate(questions["Short"], 1):
            q_text = q if isinstance(q, str) else str(q)
            doc.add_paragraph(f"{i}. {q_text}")
            doc.add_paragraph()

    # Handle Long Questions
    if "Long" in questions and questions["Long"]:
        doc.add_heading("Q.3 - Long Questions", level=1)
        for i, q in enumerate(questions["Long"], 1):
            q_text = q if isinstance(q, str) else str(q)
            doc.add_paragraph(f"{i}. {q_text}")
            doc.add_paragraph()

    # Save document
    os.makedirs("output/PU", exist_ok=True)
    output_path = f"output/PU/Generated_Paper.docx"
    doc.save(output_path)
    return doc, output_path


token=env.TOKEN
endpoint=env.ENDPOINT

embedding_model = OpenAIEmbeddings(
    base_url=endpoint,
    api_key=token,
    model="text-embedding-3-small"
)


def get_collection(teacher_id: str, filename: str):
    # Path per teacher + per file
    collection_path = os.path.join("chroma_db", teacher_id, filename)
    os.makedirs(collection_path, exist_ok=True)

    # A client scoped to THIS file’s folder
    teacher_client = chromadb.PersistentClient(path=collection_path)

    return Chroma(
        client=teacher_client,
        collection_name="index",           # always "index" inside file folder
        embedding_function=embedding_model,
        persist_directory=collection_path
    )


      

@router.post("/generate-paper")
def generate_paper(
    teacher_id: str = Form(...),
    topic_names: List[str] = Form(...),
    university: str = Form(...),
    subject: str = Form(...),
    department: str = Form(...),
    semester: str = Form(...),
    term: str = Form(...),
    dflevelMCQ: str = Form(...),
    dflevelShort: str = Form(...),
    dflevelLong: str = Form(...),
    total_marks: int = Form(...),
    duration_minutes: int = Form(...),
    questions_count: str = Form(...),
    syllabus_file_ids: List[str] = Form(...),
):
    downloaded_files = []

    # === Step 1: Download syllabus files from Node backend
    for file_id in syllabus_file_ids:
        try:
            meta_response = requests.get(f"{env.API_URI}/syllabusRoutes/getSyllabusByID/{file_id}")
            print("✅ meta_response",meta_response)
            if meta_response.status_code != 200:
                print(f"❌ Metadata fetch failed for file ID: {file_id}")
                continue

            file_meta = meta_response.json()
            file_data = file_meta.get("syllabus")
            if not file_data:
                print(f"❌ No 'syllabus' key in metadata for file ID: {file_id}")
                continue

            file_url = f"{env.API_URI}{file_data['file']}"
            filename = file_data['filename']

            file_response = requests.get(file_url)
            if file_response.status_code != 200:
                print(f"❌ Failed to download file: {file_url}")
                continue

            file_save_path = os.path.join("temp_files", filename)
            os.makedirs("temp_files", exist_ok=True)

            with open(file_save_path, "wb") as f:
                f.write(file_response.content)

            downloaded_files.append(file_save_path)


        except Exception as e:
            print(f"❌ Error processing file ID {file_id}: {e}")

    print("✅ Downloaded files:", downloaded_files)

    # === Step 2: Setup Langchain
    try:

        def get_file_hash(path):
            with open(path, "rb") as f:
                return hashlib.sha256(f.read()).hexdigest()

        def is_text_based_pdf(path, min_chars=200):
            try:
                reader = PdfReader(path)
                total_text = "".join(p.extract_text() or "" for p in reader.pages)
                return len(total_text.strip()) >= min_chars
            except:
                return False

        # === Step 3: Process downloaded files
        base_path = f"chroma_db/{teacher_id}"
        os.makedirs(base_path, exist_ok=True)
        
        for file_path in downloaded_files:
            filename = os.path.basename(file_path)
            print("✅ Processing file:", filename)
            file_hash = get_file_hash(file_path)
            file_db_path = os.path.join(base_path, filename)
            os.makedirs(file_db_path, exist_ok=True)
            hash_file = os.path.join(file_db_path, "hash.json")
        
            # Skip reprocessing if hash matches
            if os.path.exists(hash_file):
                with open(hash_file) as f:
                    stored = json.load(f)
                if stored.get("hash") == file_hash:
                    print(f"⏩ Skipping {filename}, hash unchanged")
                    os.remove(file_path)
                    continue
        
            docs = []
            try:
                if filename.endswith(".pdf") and is_text_based_pdf(file_path):
                    loader = PyPDFLoader(file_path)
                    docs = loader.load()
                elif filename.endswith(".txt"):
                    for enc in ["utf-8", "utf-16", "latin-1"]:
                        try:
                            loader = TextLoader(file_path, encoding=enc)
                            docs = loader.load()
                            if docs:
                                break
                        except:
                            pass
                elif filename.endswith(".docx"):
                    loader = Docx2txtLoader(file_path)
                    docs = loader.load()
            except Exception as e:
                print(f"⚠️ Failed to load file {filename}: {e}")
        
            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            chunks = [c for c in splitter.split_documents(docs) if c.page_content.strip()]
        
            if chunks:
                print(f"✅ Loaded {len(chunks)} new chunks. Inserting into Chroma...")
                collection = get_collection(teacher_id, filename)
                
                collection.add_documents(chunks)
                print("✅ before loading chunks")
                with open(hash_file, "w") as f:
                    json.dump({"hash": file_hash}, f)
        
            os.remove(file_path)
        
        
        # === Step 4: Query for topics (outside loop)
        all_chunks = []
        used_files = set()
        unmatched_topics = []
        
        for topic_query in topic_names:
            matched_chunks = []
            keywords = topic_query.lower().split()
        
            for folder in os.listdir(base_path):
                collection = get_collection(teacher_id, folder)
        
                results = collection.similarity_search(topic_query, k=10)
        
                topic_matched_chunks = [
                    doc.page_content for doc in results
                    if topic_query.lower() in doc.page_content.lower()
                    or all(word in doc.page_content.lower() for word in keywords)
                ]
        
                if topic_matched_chunks:
                    matched_chunks.extend(topic_matched_chunks)
                    used_files.add(folder)
        
            if matched_chunks:
                all_chunks.extend(matched_chunks)
            else:
                unmatched_topics.append(topic_query)
        
        
        print("all_chunks", all_chunks)

    
        # ✅ Throw error if nothing matched at all
        if not all_chunks:
            return {
                "error": f"❌ None of the given topics were found in any uploaded file.",
                "unmatched_topics": unmatched_topics
            }
    
        # ✅ If some matched and some didn’t, return the unmatched for info
        if unmatched_topics:
            print(f"⚠️ These topics were not found: {unmatched_topics}")

        syllabus_text = "\n".join(chunk.strip() for chunk in all_chunks)
        count_dict = json.loads(questions_count)
        print("syllabus_text",syllabus_text)
        data = PaperRequest(
            university=university,
            subject=subject,
            department=department,
            semester=semester,
            term=term,
            dflevelMCQ=dflevelMCQ,
            dflevelShort=dflevelShort,
            dflevelLong=dflevelLong,
            total_marks=total_marks,
            duration_minutes=duration_minutes,
            questions_count=count_dict,
        )

        questions = generate_exam_paper(syllabus_text, count_dict, data, topic_names)
        doc, output_path = format_paper_pu(questions, data)


        node_upload_url = f"{env.API_URI}/paperRoutes/save-paper"

        upload_data = {
            'teacher_id': teacher_id,
            'subject': subject,
            'department': department,
            'university': university,
            'semester': semester,
            'term': term,
            'dflevelMCQ': dflevelMCQ,
            'dflevelShort': dflevelShort,
            'dflevelLong': dflevelLong,
            'total_marks': str(total_marks),
            'duration_minutes': str(duration_minutes),
            'questions_count': questions_count,
        }

        for topic in topic_names:
            upload_data.setdefault("topic_names", []).append(topic)

        for sid in syllabus_file_ids:
            upload_data.setdefault("syllabus_file_ids", []).append(sid)

        try:
            with open(output_path, "rb") as f:
                files = {
                    "generated_paper": (
                        os.path.basename(output_path),
                        f,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                }
                node_response = requests.post(node_upload_url, data=upload_data, files=files)
                node_response.raise_for_status()
                paperId = node_response.json().get("paperId")
        except Exception as e:
            return {"error": f"Failed to upload to Node.js backend: {e}"}

        # ✅ Delete AFTER file is closed (outside `with`)
        if os.path.exists(output_path):
            os.remove(output_path)

            
        return {
            "paper_generated": True,
            "used_files": list(used_files),
            "unmatched_topics": unmatched_topics,
            "paperId": paperId
        }

    except Exception as e:
        return {"error": str(e)}
